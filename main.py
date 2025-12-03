from fastapi import FastAPI, UploadFile, File, Form, Request
from fastapi.responses import JSONResponse, Response, HTMLResponse
from fastapi.templating import Jinja2Templates
from docx import Document
from io import BytesIO
import re
from pathlib import Path
import json
import traceback # Importante para ver el error real en consola

app = FastAPI()

# Config directorios
current_dir = Path(__file__).parent.resolve()
templates = Jinja2Templates(directory=str(current_dir / "templates"))

# Regex para campos manuales {{campo}} (sin '!')
campo_regex = re.compile(r"\{\{\s*([^\{\}!]+?)\s*\}\}")

def extraer_campos_de_parrafos(parrafos):
    campos = set()
    for p in parrafos:
        # Manejo de error si p.text es None
        if p.text:
            for match in campo_regex.finditer(p.text):
                campos.add(match.group(1).strip())
    return campos

def extraer_campos_de_tablas(tablas):
    campos = set()
    for table in tablas:
        for row in table.rows:
            for cell in row.cells:
                campos |= extraer_campos_de_parrafos(cell.paragraphs)
    return campos

def extraer_todos_los_campos(doc: Document):
    campos = set()
    campos |= extraer_campos_de_parrafos(doc.paragraphs)
    campos |= extraer_campos_de_tablas(doc.tables)

    for section in doc.sections:
        try:
            campos |= extraer_campos_de_parrafos(section.header.paragraphs)
            campos |= extraer_campos_de_tablas(section.header.tables)
            campos |= extraer_campos_de_parrafos(section.footer.paragraphs)
            campos |= extraer_campos_de_tablas(section.footer.tables)

            if section.different_first_page_header_footer:
                campos |= extraer_campos_de_parrafos(section.first_page_header.paragraphs)
                campos |= extraer_campos_de_tablas(section.first_page_header.tables)
                campos |= extraer_campos_de_parrafos(section.first_page_footer.paragraphs)
                campos |= extraer_campos_de_tablas(section.first_page_footer.tables)
        except Exception:
            # Algunas secciones pueden no existir o dar error al acceder
            pass
            
    return sorted(campos)

def reemplazar_texto_en_parrafo(parrafo, reemplazos):
    if not parrafo.runs:
        return
    texto_completo = "".join(run.text for run in parrafo.runs)
    
    def reemplazo_match(match):
        campo = match.group(1).strip()
        return str(reemplazos.get(campo, match.group(0)))
    
    nuevo_texto = campo_regex.sub(reemplazo_match, texto_completo)

    if nuevo_texto != texto_completo:
        parrafo.runs[0].text = nuevo_texto
        for run in parrafo.runs[1:]:
            run.text = ""

def reemplazar_campos(doc: Document, reemplazos: dict):
    for p in doc.paragraphs:
        reemplazar_texto_en_parrafo(p, reemplazos)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    reemplazar_texto_en_parrafo(p, reemplazos)

    for section in doc.sections:
        try:
            for p in section.header.paragraphs:
                reemplazar_texto_en_parrafo(p, reemplazos)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            reemplazar_texto_en_parrafo(p, reemplazos)

            for p in section.footer.paragraphs:
                reemplazar_texto_en_parrafo(p, reemplazos)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            reemplazar_texto_en_parrafo(p, reemplazos)
            
            if section.different_first_page_header_footer:
                for p in section.first_page_header.paragraphs:
                    reemplazar_texto_en_parrafo(p, reemplazos)
                for table in section.first_page_header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                reemplazar_texto_en_parrafo(p, reemplazos)

                for p in section.first_page_footer.paragraphs:
                    reemplazar_texto_en_parrafo(p, reemplazos)
                for table in section.first_page_footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                reemplazar_texto_en_parrafo(p, reemplazos)
        except Exception:
            pass

@app.post("/detectar-campos")
async def detectar_campos(archivo_word: UploadFile = File(...)):
    try:
        if not archivo_word.filename.endswith(".docx"):
            return JSONResponse(status_code=400, content={"detail": "Solo se aceptan archivos .docx"})
            
        contenido = await archivo_word.read()
        doc = Document(BytesIO(contenido))
        campos = extraer_todos_los_campos(doc)
        return JSONResponse({"campos": campos})
        
    except Exception as e:
        # IMPRIMIR EL ERROR REAL EN LA CONSOLA
        print("----------- ERROR EN DETECTAR CAMPOS -----------")
        traceback.print_exc() 
        print("------------------------------------------------")
        # Devolver JSON forzado
        return JSONResponse(status_code=500, content={"detail": f"Error interno: {str(e)}"})

@app.post("/procesar-manual")
async def procesar_manual(
    archivo_word: UploadFile = File(...),
    replacements: str = Form(...)
):
    try:
        nombre_original = archivo_word.filename
        reemplazos = json.loads(replacements)
        contenido = await archivo_word.read()
        doc = Document(BytesIO(contenido))
        reemplazar_campos(doc, reemplazos)
        salida = BytesIO()
        doc.save(salida)
        salida.seek(0)
        
        return Response(
            content=salida.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{nombre_original}"'}
        )
    except Exception as e:
        print("----------- ERROR EN PROCESAR MANUAL -----------")
        traceback.print_exc()
        print("------------------------------------------------")
        return JSONResponse(status_code=500, content={"detail": f"Error interno: {str(e)}"})

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

# NO AGREGUES EL EXCEPTION HANDLER AQUI
